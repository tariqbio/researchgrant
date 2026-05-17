import os
import re
import uuid
from urllib.request import Request, urlopen
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form, BackgroundTasks
from sqlalchemy.orm import Session
from typing import Optional

from app.db.session import get_db
from app.core.security import get_current_admin, get_current_user
from app.core.config import settings
from app.models.pipeline import IngestionJob, Source, CommunitySubmission
from app.schemas.pipeline import (
    IngestionJobOut, CommunitySubmissionCreate,
    CommunitySubmissionOut, SourceOut, SourceCreate,
)

router = APIRouter(prefix="/pipeline", tags=["pipeline"])


def save_upload(file_bytes: bytes, filename: str) -> str:
    """Save uploaded file to local storage. Swap for S3 in production."""
    os.makedirs(settings.STORAGE_LOCAL_PATH, exist_ok=True)
    unique_name = f"{uuid.uuid4()}_{filename}"
    path = os.path.join(settings.STORAGE_LOCAL_PATH, unique_name)
    with open(path, "wb") as f:
        f.write(file_bytes)
    return path


# ── Upload & pipeline ─────────────────────────────────────────────────────────

@router.post("/upload", response_model=IngestionJobOut)
async def upload_pdf(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    source_id: Optional[str] = Form(None),
    db: Session = Depends(get_db),
    admin=Depends(get_current_admin),
):
    """Admin uploads a PDF/DOCX/TXT grant notice. Creates an IngestionJob and queues pipeline."""
    filename = file.filename or "uploaded_notice"
    if not filename.lower().endswith((".pdf", ".docx", ".txt")):
        raise HTTPException(status_code=400, detail="Only PDF, DOCX, or TXT files are accepted")

    file_bytes = await file.read()
    if len(file_bytes) > 20 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="File too large (max 20MB)")

    file_path = save_upload(file_bytes, filename)

    job = IngestionJob(
        source_id=uuid.UUID(source_id) if source_id else None,
        raw_file_path=file_path,
        job_status="pending_ocr",
    )
    db.add(job)
    db.commit()
    db.refresh(job)

    # Use FastAPI BackgroundTasks — safe inside a request context
    background_tasks.add_task(run_pipeline, job.id)

    return job


@router.get("/jobs/{job_id}", response_model=IngestionJobOut)
def get_job_status(
    job_id: uuid.UUID,
    db: Session = Depends(get_db),
    admin=Depends(get_current_admin),
):
    """Poll pipeline job status. Frontend polls this every 3s during upload."""
    job = db.query(IngestionJob).filter(IngestionJob.id == job_id).first()
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    return job


@router.get("/jobs", response_model=list[IngestionJobOut])
def list_jobs(
    limit: int = 50,
    db: Session = Depends(get_db),
    admin=Depends(get_current_admin),
):
    """List recent ingestion jobs for the admin dashboard."""
    return (
        db.query(IngestionJob)
        .order_by(IngestionJob.created_at.desc())
        .limit(limit)
        .all()
    )


async def run_pipeline(job_id: uuid.UUID):
    """
    Background pipeline runner called via BackgroundTasks.
    Step 1: OCR (Google Vision / PyMuPDF fallback)
    Step 2: AI extraction (Claude API)
    Step 3: Create draft Grant → status = pending_review
    """
    from app.db.session import SessionLocal
    from app.services.extraction_service import extract_grant_from_text, build_grant_from_extraction

    db = SessionLocal()
    try:
        job = db.query(IngestionJob).filter(IngestionJob.id == job_id).first()
        if not job:
            return

        # Step 1 — OCR
        job.job_status = "ocr_running"
        db.commit()

        try:
            raw_text, engine, confidence = extract_text_from_source(job.raw_file_path, job.raw_url)
            job.ocr_engine = engine
            job.ocr_confidence = confidence
            job.raw_text_path = save_raw_text(raw_text, job.id)
            job.job_status = "ai_running"
            db.commit()
        except Exception as e:
            job.job_status = "ocr_failed"
            job.failure_reason = str(e)
            db.commit()
            return

        if not raw_text.strip():
            job.job_status = "ocr_failed"
            job.failure_reason = "OCR returned empty text — document may be blank or unreadable"
            db.commit()
            return

        # Step 2 — AI extraction
        try:
            extracted = await extract_grant_from_text(raw_text)
            job.ai_extracted_json = extracted
            job.ai_model = extracted.get("_ai_model") or "claude-sonnet-4-20250514"
            job.job_status = "pending_review"
            db.commit()
        except Exception as e:
            job.job_status = "ai_failed"
            job.failure_reason = str(e)
            db.commit()
            return

        # Step 3 — Create draft Grant linked to this job
        from app.models.grant import Grant

        grant_data = build_grant_from_extraction(extracted)
        grant = Grant(
            **grant_data,
            ingestion_job_id=job.id,
            source_id=job.source_id,
            original_pdf_path=job.raw_file_path,
            status="pending_review",
        )
        db.add(grant)
        db.commit()

    except Exception as e:
        # Catch-all so background task never silently crashes
        try:
            job = db.query(IngestionJob).filter(IngestionJob.id == job_id).first()
            if job:
                job.job_status = "ai_failed"
                job.failure_reason = f"Unhandled error: {str(e)}"
                db.commit()
        except Exception:
            pass
    finally:
        db.close()


def save_raw_text(raw_text: str, job_id: uuid.UUID) -> str:
    os.makedirs(settings.STORAGE_LOCAL_PATH, exist_ok=True)
    path = os.path.join(settings.STORAGE_LOCAL_PATH, f"{job_id}_ocr.txt")
    with open(path, "w", encoding="utf-8") as f:
        f.write(raw_text)
    return path


def extract_text_from_source(file_path: Optional[str], raw_url: Optional[str]) -> tuple[str, str, float]:
    if raw_url and not file_path:
        return extract_text_from_url(raw_url), "url_text", 0.55
    if not file_path:
        raise ValueError("No uploaded file or URL was attached to this job")

    lower = file_path.lower()
    if lower.endswith(".pdf"):
        text = extract_text_from_pdf(file_path)
        if text.strip():
            return text, "pymupdf_text", 0.9
        text = google_vision_pdf_ocr(file_path)
        if text.strip():
            return text, "google_vision", 0.75
        return "", "pymupdf_empty", 0.0

    if lower.endswith(".docx"):
        return extract_text_from_docx(file_path), "python_docx", 0.9

    if lower.endswith(".txt"):
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read(), "plain_text", 1.0

    raise ValueError("Unsupported uploaded file type")


def extract_text_from_pdf(file_path: str) -> str:
    import fitz

    doc = fitz.open(file_path)
    try:
        return "\n".join(page.get_text() for page in doc)
    finally:
        doc.close()


def google_vision_pdf_ocr(file_path: str) -> str:
    if not (settings.GOOGLE_APPLICATION_CREDENTIALS or settings.GOOGLE_VISION_CREDENTIALS):
        return ""
    try:
        import fitz
        from google.cloud import vision
    except Exception:
        return ""

    if settings.GOOGLE_VISION_CREDENTIALS and not os.environ.get("GOOGLE_APPLICATION_CREDENTIALS"):
        os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = settings.GOOGLE_VISION_CREDENTIALS

    client = vision.ImageAnnotatorClient()
    doc = fitz.open(file_path)
    chunks: list[str] = []
    try:
        for page_index in range(min(8, len(doc))):
            page = doc[page_index]
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
            image = vision.Image(content=pix.tobytes("png"))
            response = client.document_text_detection(image=image)
            if not response.error.message:
                chunks.append(response.full_text_annotation.text or "")
    finally:
        doc.close()
    return "\n".join(chunks)


def extract_text_from_docx(file_path: str) -> str:
    from docx import Document

    doc = Document(file_path)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract_text_from_url(raw_url: str) -> str:
    req = Request(raw_url, headers={"User-Agent": "GrantBD/1.0"})
    with urlopen(req, timeout=12) as response:
        html = response.read(2_000_000).decode("utf-8", errors="ignore")
    html = re.sub(r"(?is)<(script|style).*?>.*?</\1>", " ", html)
    text = re.sub(r"(?s)<[^>]+>", " ", html)
    return re.sub(r"\s+", " ", text).strip()


# ── Community submissions ─────────────────────────────────────────────────────

@router.post("/submit", response_model=CommunitySubmissionOut, status_code=201)
def submit_grant_url(
    payload: CommunitySubmissionCreate,
    db: Session = Depends(get_db),
    current_user=Depends(get_current_user),
):
    """Any logged-in researcher can submit a grant URL they found."""
    submission = CommunitySubmission(
        submitted_by=current_user.id,
        source_url=payload.source_url,
        notes=payload.notes,
    )
    db.add(submission)
    db.commit()
    db.refresh(submission)
    return submission


@router.get("/submissions", response_model=list[CommunitySubmissionOut])
def list_submissions(
    db: Session = Depends(get_db),
    admin=Depends(get_current_admin),
):
    return (
        db.query(CommunitySubmission)
        .filter(CommunitySubmission.status == "pending")
        .order_by(CommunitySubmission.created_at.desc())
        .all()
    )


@router.post("/submissions/{submission_id}/approve")
def approve_submission(
    submission_id: uuid.UUID,
    background_tasks: BackgroundTasks,
    db: Session = Depends(get_db),
    admin=Depends(get_current_admin),
):
    """
    Admin approves a community submission URL.
    Fetches the page, saves as PDF placeholder, queues pipeline.
    """
    sub = db.query(CommunitySubmission).filter(CommunitySubmission.id == submission_id).first()
    if not sub:
        raise HTTPException(status_code=404, detail="Submission not found")

    # Create job with source URL instead of a file path
    job = IngestionJob(
        job_status="pending_ocr",
        raw_url=sub.source_url,
    )
    db.add(job)
    sub.status = "approved"
    db.commit()
    db.refresh(job)

    background_tasks.add_task(run_pipeline, job.id)
    return {"job_id": str(job.id), "status": "queued"}


@router.post("/submissions/{submission_id}/reject")
def reject_submission(
    submission_id: uuid.UUID,
    db: Session = Depends(get_db),
    admin=Depends(get_current_admin),
):
    sub = db.query(CommunitySubmission).filter(CommunitySubmission.id == submission_id).first()
    if not sub:
        raise HTTPException(status_code=404, detail="Submission not found")
    sub.status = "rejected"
    db.commit()
    return {"status": "rejected"}


# ── Sources management ────────────────────────────────────────────────────────

@router.post("/sources", response_model=SourceOut, status_code=201)
def create_source(
    payload: SourceCreate,
    db: Session = Depends(get_db),
    admin=Depends(get_current_admin),
):
    source = Source(**payload.model_dump())
    db.add(source)
    db.commit()
    db.refresh(source)
    return source


@router.get("/sources", response_model=list[SourceOut])
def list_sources(db: Session = Depends(get_db)):
    """Public — used by the frontend to populate source filters."""
    return db.query(Source).order_by(Source.name).all()
